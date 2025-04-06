import requests
from typing import List, Dict, Any, Union, Optional
import json
from pypdf import PdfReader
import pandas as pd
from datetime import datetime, timedelta
import chardet
import os
import hashlib
import re
import csv
import tempfile
# LangChain Imports
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field, field_validator
from typing_extensions import TypedDict

# Initialize with proper type annotations for Python 3.8+ compatibility
from langchain_community.chat_models import ChatOllama

# Define Pydantic models for structured output
class JobRequirements(BaseModel):
    title: Optional[str] = Field(description="The specific position title")
    company: Optional[str] = Field(description="The hiring company")
    required_skills: Optional[List[str]] = Field(description="List of technical and soft skills required")
    experience_years: Optional[int] = Field(description="Required years of experience in the field")
    qualifications: Optional[List[str]] = Field(description="Required educational qualifications")
    responsibilities: Optional[List[str]] = Field(description="Key job responsibilities and duties")

class EducationEntry(BaseModel):
    degree: Optional[str] = Field(description="Specific degree name")
    institution: Optional[str] = Field(description="Institution name")
    year: Optional[str] = Field(description="Completion year")

class ExperienceEntry(BaseModel):
    title: Optional[str] = Field(description="Job title")
    company: Optional[str] = Field(description="Company name")
    duration: Optional[str] = Field(description="Employment duration")
    start_date: Optional[str] = Field(description="Start date (YYYY-MM-DD format)")
    end_date: Optional[str] = Field(description="End date (YYYY-MM-DD format) or 'Present' if current position")
    years: Optional[float] = Field(description="Total years in this role (calculated)")
    description: Optional[str] = Field(description="Key responsibilities and achievements")

class CandidateInfo(BaseModel):
    name: Optional[str] = Field(description="Candidate's full name")
    email: Optional[str] = Field(description="Candidate's email address")
    phone: Optional[str] = Field(description="Candidate's phone number")
    education: Optional[List[EducationEntry]] = Field(description="List of education entries")
    experience: Optional[List[ExperienceEntry]] = Field(description="List of work experience entries")
    skills: Optional[List[str]] = Field(description="List of skills and technical competencies")

class InterviewRequest(BaseModel):
    subject: str = Field(description="Clear and engaging email subject line")
    body: str = Field(description="Personalized and professional email body text")
    proposed_dates: List[str] = Field(description="List of proposed interview dates (YYYY-MM-DD)")
    interview_type: str = Field(description="Interview type (e.g., online, in-person)")

class JobDescriptionAgent:
    """Agent responsible for parsing job descriptions from CSV files."""
    
    def __init__(self):
        """Initialize the JobDescriptionAgent."""
        base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        self.llm = ChatOllama(
            model="llama3",
            base_url=base_url,
            format="json",
            temperature=0.2
        )
    
    def parse_jd_csv(self, filepath: str) -> List[Dict[str, str]]:
        """
        Parse a CSV file containing job descriptions.
        
        Args:
            filepath: Path to the CSV file
            
        Returns:
            List of dictionaries with job information
        """
        print(f"Attempting to read CSV file: {filepath}")
        
        # Try different encodings
        encodings_to_try = ['utf-8', 'latin-1', 'ISO-8859-1']
        csv_data = None
        used_encoding = None
        
        for encoding in encodings_to_try:
            try:
                print(f"Trying encoding: {encoding}")
                with open(filepath, 'r', encoding=encoding) as f:
                    csv_content = f.read()
                    if csv_content:
                        # Successfully read with this encoding
                        used_encoding = encoding
                        break
            except Exception as e:
                print(f"Failed with encoding {encoding}: {e}")
        
        if not used_encoding:
            # Try to detect encoding with chardet
            try:
                with open(filepath, 'rb') as f:
                    raw_data = f.read()
                    result = chardet.detect(raw_data)
                    detected_encoding = result['encoding']
                    print(f"Detected encoding: {detected_encoding}")
                    with open(filepath, 'r', encoding=detected_encoding) as f:
                        csv_content = f.read()
                        if csv_content:
                            used_encoding = detected_encoding
            except Exception as e:
                print(f"Failed to detect encoding: {e}")
                return []
        
        if not used_encoding:
            print("Could not read file with any encoding")
            return []
            
        print(f"Successfully read CSV with encoding: {used_encoding}")
        
        # Process the CSV content
        try:
            reader = csv.DictReader(csv_content.splitlines())
            headers = reader.fieldnames
            print(f"CSV columns: {headers}")
            
            # Check required columns
            if not headers or not any(header in headers for header in ['Job Title', 'Description', 'Job Description']):
                print("CSV missing required columns (Job Title and Description/Job Description)")
                return []
                
            # Map the headers to standardized keys
            header_mapping = {}
            for h in headers:
                if h.lower() in ['job title', 'title', 'position']:
                    header_mapping[h] = 'title'
                elif h.lower() in ['job description', 'description']:
                    header_mapping[h] = 'description'
                else:
                    # Keep other headers as is
                    header_mapping[h] = h
            
            processed_records = []
            
            for i, row in enumerate(reader, 1):
                processed_row = {}
                
                # Map the values using the header mapping
                for original_header, value in row.items():
                    if original_header in header_mapping:
                        processed_row[header_mapping[original_header]] = value
                
                # Ensure required fields exist
                if 'title' not in processed_row or not processed_row['title']:
                    print(f"Record {i} missing title, skipping")
                    continue
                    
                if 'description' not in processed_row or not processed_row['description']:
                    print(f"Record {i} missing description, skipping")
                    continue
                
                print(f"Processing record {i}:")
                print(f"Job Title: {processed_row['title']}")
                
                processed_records.append(processed_row)
                print(f"Processed record title: {processed_row['title']}")
            
            print(f"Total processed records: {len(processed_records)}")
            return processed_records
            
        except Exception as e:
            print(f"Error processing CSV: {e}")
            return []
    
    def extract_key_requirements(self, job_description_text: str) -> Dict[str, Any]:
        """Extract key requirements from job description text."""
        print(f"Extracting requirements from description (length: {len(job_description_text)} chars)")
        prompt = f"""
        Analyze the following job description text and extract the key requirements.
        
        Job Description Text:
        {job_description_text}
        
        Extract the following information and return ONLY valid JSON:
        - job_title: The title of the position
        - company_name: The name of the company (if available)
        - description: A brief summary of the job role
        - required_skills: List of mandatory technical skills. 
        - preferred_skills: List of desired technical skills.
        - qualifications: Required education level (e.g., Bachelor's, Master's), certifications.
        - responsibilities: List of key duties and tasks.
        - required_experience_num: The minimum number of years required (numeric). Use 0 if not specified.
        
        IMPORTANT INSTRUCTIONS FOR SKILLS:
        1. Identify skills explicitly listed under headings like "Skills", "Requirements", "Qualifications".
        2. **Also, carefully read the 'responsibilities' and 'description' sections. Extract any technical skills mentioned within these sections.**
        3. Classify extracted skills as 'required' if the text implies necessity (e.g., "must have", "required", "experience in X needed").
        4. Classify extracted skills as 'preferred' if the text implies desirability (e.g., "nice to have", "bonus", "familiarity with Y").
        5. If uncertain, classify skills found in responsibilities/description as 'preferred'.
        6. Be specific (e.g., "Python", "AWS", "React Native").

        Return ONLY valid JSON.
        """
        
        extracted_requirements = {}
        try:
            # Define the output structure using Pydantic or TypedDict
            class Requirements(TypedDict):
                required_skills: List[str]
                preferred_skills: List[str]
                experience_years: Union[int, str] # Allow string initially for parsing robustness
                education_level: str

            # Create a parser for the structured output
            parser = PydanticOutputParser(pydantic_object=Requirements)

            # Create a prompt template
            prompt_template = ChatPromptTemplate.from_messages(
                [SystemMessage(content="You are an expert HR assistant specializing in extracting structured job requirements."),
                 HumanMessage(content=prompt)])
            
            # Create the chain
            chain = prompt_template | self.llm | parser
            
            # Invoke the chain
            response = chain.invoke({})
            extracted_requirements = response

            # --- Post-processing for experience_years --- 
            raw_experience = extracted_requirements.get('experience_years', 0)
            processed_years = 0 # Default to 0
            try:
                if isinstance(raw_experience, int):
                    processed_years = raw_experience
                elif isinstance(raw_experience, str):
                    # Check if the string is empty or "None"
                    if not raw_experience or raw_experience.lower() == 'none':
                        # Still default to 0, but with better logging
                        print(f"No specific experience requirement found, defaulting to 0 years")
                        processed_years = 0
                    else:
                        # Try to find numbers like '3+', '5', '7-10'
                        matches = re.findall(r'\d+\.\d+|\d+', raw_experience)
                        if matches:
                            # Take the first number found as the minimum years
                            processed_years = int(float(matches[0]))
                            print(f"Extracted {processed_years} years from '{raw_experience}'")
                        else:
                            # Handle more complex experience descriptions without numbers
                            text_lower = raw_experience.lower()
                            
                            # Check for keywords indicating specific experience levels
                            if any(senior_term in text_lower for senior_term in ['senior', 'extensive', 'substantial']):
                                processed_years = 5
                                print(f"Inferred 5 years from senior-level description: '{raw_experience}'")
                            elif any(mid_term in text_lower for mid_term in ['mid', 'intermediate', 'moderate']):
                                processed_years = 3
                                print(f"Inferred 3 years from mid-level description: '{raw_experience}'")
                            elif any(entry_term in text_lower for entry_term in ['entry', 'junior', 'beginner']):
                                processed_years = 1
                                print(f"Inferred 1 year from entry-level description: '{raw_experience}'")
                            elif any(exp_term in text_lower for exp_term in ['proven', 'demonstrated', 'significant']):
                                processed_years = 4  # Assume "proven experience" means ~4 years
                                print(f"Inferred 4 years from 'proven experience' description: '{raw_experience}'")
                            elif any(exp_term in text_lower for exp_term in ['some', 'basic', 'minimum']):
                                processed_years = 2  # Assume "some experience" means ~2 years
                                print(f"Inferred 2 years from 'some experience' description: '{raw_experience}'")
                            else:
                                # Fallback: if we detect "experience" but can't quantify it
                                if 'experience' in text_lower:
                                    processed_years = 3  # Reasonable default
                                    print(f"Unquantified experience requirement, assuming 3 years: '{raw_experience}'")
            except (ValueError, TypeError):
                print(f"Warning: Could not reliably parse experience_years: '{raw_experience}', defaulting to 0.")
                processed_years = 0 # Default to 0 if parsing fails
            
            # Ensure it's non-negative
            processed_years = max(0, processed_years)
            extracted_requirements['experience_years'] = processed_years
            # --- End post-processing --- 

            # Add the raw JSON for storage if needed
            extracted_requirements['requirements_json'] = json.dumps(extracted_requirements)
            print(f"Successfully extracted requirements: {extracted_requirements}")

        except Exception as e:
            print(f"Error during requirement extraction: {e}")
            # Provide default structure on error
            extracted_requirements = { 
                "required_skills": [],
                "preferred_skills": [],
                "experience_years": 0,
                "education_level": "None",
                "requirements_json": "{}"
            }
            
        return extracted_requirements

class CandidateAgent:
    """Agent responsible for parsing candidate CVs."""
    
    def __init__(self):
        """Initialize the CandidateAgent."""
        base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        self.llm = ChatOllama(
            model="llama3",
            base_url=base_url,
            format="json",
            temperature=0.2
        )
    
    def parse_cv(self, filepath: str) -> str:
        """
        Parse a CV and extract the text content using appropriate extraction methods.
        
        Args:
            filepath: Path to the CV file
            
        Returns:
            Extracted text from the CV
        """
        print(f"Parsing CV: {filepath}")
        
        try:
            # Check if file exists
            if not os.path.exists(filepath):
                print(f"File not found: {filepath}")
                return ""
                
            # Read the file as binary
            with open(filepath, 'rb') as file:
                file_content = file.read()
                
            # Get the file extension
            _, ext = os.path.splitext(filepath)
            ext = ext.lower()
            
            # For text-based files, try to read directly
            if ext in ['.txt', '.md', '.html', '.htm']:
                try:
                    # Try different encodings
                    encodings = ['utf-8', 'latin-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            return file_content.decode(encoding)
                        except UnicodeDecodeError:
                            continue
                except Exception as e:
                    print(f"Error decoding text file: {e}")
            
            # For PDF files, use pdfminer.six to extract text
            elif ext == '.pdf':
                try:
                    # Import pdfminer.six components here to avoid dependency issues
                    from pdfminer.high_level import extract_text
                    from pdfminer.pdfparser import PDFSyntaxError
                    
                    # Check if it's a valid PDF file
                    if not file_content.startswith(b'%PDF-'):
                        print("Not a valid PDF file")
                        return ""
                    
                    # Extract text from the PDF
                    try:
                        extracted_text = extract_text(filepath)
                        if extracted_text:
                            print(f"Successfully extracted {len(extracted_text)} characters from PDF")
                            return extracted_text
                        else:
                            print("PDF extraction returned empty text")
                    except PDFSyntaxError as e:
                        print(f"PDF syntax error: {e}")
                        return ""
                        
                    # If pdfminer extraction fails, try fallback method
                    print("Primary PDF extraction failed, trying fallback method...")
                    from io import StringIO
                    from pdfminer.layout import LAParams
                    from pdfminer.converter import TextConverter
                    from pdfminer.pdfdocument import PDFDocument
                    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
                    from pdfminer.pdfpage import PDFPage
                    
                    output = StringIO()
                    with open(filepath, 'rb') as in_file:
                        parser = PDFDocument(PDFParser(in_file))
                        rsrcmgr = PDFResourceManager()
                        device = TextConverter(rsrcmgr, output, laparams=LAParams())
                        interpreter = PDFPageInterpreter(rsrcmgr, device)
                        for page in PDFPage.create_pages(parser):
                            interpreter.process_page(page)
                    
                    text = output.getvalue()
                    output.close()
                    
                    if text:
                        print(f"Fallback PDF extraction successful: {len(text)} characters")
                        return text
                    else:
                        print("Fallback PDF extraction also failed")
                        return ""
                        
                except ImportError as e:
                    print(f"PDF extraction libraries not available: {e}")
                    print("Please install pdfminer.six with: pip install pdfminer.six")
                    return f"[PDF extraction failed. Required libraries not available: {e}]"
                except Exception as e:
                    print(f"Error with PDF extraction: {e}")
                    return ""
            
            # For doc/docx, try to use python-docx if available
            elif ext in ['.doc', '.docx']:
                try:
                    if ext == '.docx':
                        import docx
                        doc = docx.Document(filepath)
                        full_text = []
                        for para in doc.paragraphs:
                            full_text.append(para.text)
                        return '\n'.join(full_text)
                    else:
                        print("DOC format not supported directly, need additional libraries")
                        return f"[Document at {filepath} is in DOC format which requires additional libraries]"
                except ImportError:
                    print("docx library not available")
                    print("Please install python-docx with: pip install python-docx")
                    return f"[Word document extraction failed. Required libraries not available]"
                except Exception as e:
                    print(f"Error extracting Word document: {e}")
                    return ""
            
            # For other file types, return a message
            else:
                print(f"Unsupported file type: {ext}")
                return f"[Document at {filepath} has unsupported file type: {ext}. Please provide text content manually.]"
                
            # If we reach here, extraction failed
            return ""
            
        except Exception as e:
            print(f"Error parsing document: {e}")
            return ""
    
    def extract_candidate_info(self, cv_text: str) -> Dict[str, Any]:
        """
        Extract structured information from CV text.
        
        Args:
            cv_text: The text extracted from the CV
            
        Returns:
            Dictionary with structured candidate information
        """
        if not cv_text:
            return {
                "name": "Unknown Candidate",
                "email": "no-email@example.com",
                "skills": [],
                "education": [],
                "experience": []
            }
            
        print(f"Extracting info from CV text: {cv_text[:50]}...")
        
        # Create a prompt for the LLM
        prompt = f"""
        Extract key information from the following CV text.
        
        CV Text:
        {cv_text}
        
        Please extract and return ONLY the following information in JSON format:
        - name: The candidate's full name
        - email: The candidate's email address
        - phone: The candidate's phone number
        - skills: A list of the candidate's technical skills
        - education: A list of the candidate's educational background (degree, institution, year)
        - experience: A list of the candidate's work experience with the following fields:
          * title: The exact job title/position held
          * company: Company or organization name
          * start_date: When they started (YYYY-MM format)
          * end_date: When they ended (YYYY-MM format) or "Present" if current
          * years: Calculate exact years in this role (e.g., Jan 2019 to Dec 2023 is 5.0 years)
          * description: Brief description of responsibilities
        
        IMPORTANT INSTRUCTIONS FOR EXPERIENCE:
        1. Extract EXACT position titles - do not generalize or change them
        2. Calculate duration in years accurately from start_date and end_date. Consider the full duration, including start and end months.
           - Example 1: Jan 2019 - Dec 2023 => 5.0 years
           - Example 2: May 2019 - June 2023 => 4 years and 2 months => 4.17 years
           - Example 3: 2018 - 2022 => Assume Jan 2018 - Dec 2022 => 5.0 years
           - Example 4: If end_date is 'Present', calculate up to the current date.
        3. Include the calculated numeric value in the "years" field (e.g., 4.17).
        
        Return ONLY valid JSON.
        """
        
        try:
            # Get structured output from LLM
            response = self.llm.invoke([HumanMessage(content=prompt)])
            content = response.content
            print(f"Raw LLM response: {content}")
            
            # Normalize the response
            normalized = {}
            
            # Check if we got a dictionary that contains 'properties'
            if isinstance(content, dict) and 'properties' in content:
                # Extract data from 'properties'
                properties = content.get('properties', {})
                for key, value in properties.items():
                    normalized[key] = value
                print(f"Extracted from properties: {normalized}")
            # Handle various string response formats
            elif isinstance(content, str):
                # Try to extract JSON from the text
                try:
                    # Check if the response is already valid JSON
                    parsed = json.loads(content)
                    
                    # Check if parsed JSON has 'properties'
                    if 'properties' in parsed:
                        properties = parsed.get('properties', {})
                        for key, value in properties.items():
                            normalized[key] = value
                        print(f"Extracted from properties in JSON: {normalized}")
                    else:
                        normalized = parsed
                    
                except json.JSONDecodeError:
                    # Try to extract JSON from a text response
                    match = re.search(r'```(?:json)?(.*?)```', content, re.DOTALL)
                    if match:
                        try:
                            json_str = match.group(1).strip()
                            parsed = json.loads(json_str)
                            
                            # Check if parsed JSON has 'properties'
                            if 'properties' in parsed:
                                properties = parsed.get('properties', {})
                                for key, value in properties.items():
                                    normalized[key] = value
                                print(f"Extracted from properties in code block: {normalized}")
                            else:
                                normalized = parsed
                                
                        except json.JSONDecodeError:
                            pass
            else:
                # Use the content directly
                normalized = content
            
            # Extract basics from the text (fallback)
            if not normalized:
                name_match = re.search(r'name:?\s*([^\n]+)', cv_text, re.IGNORECASE)
                name = name_match.group(1).strip() if name_match else "Unknown Candidate"
                
                email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', cv_text)
                email = email_match.group(0) if email_match else "no-email@example.com"
                
                # Return basic info
                normalized = {
                    "name": name,
                    "email": email,
                    "skills": [],
                    "education": [],
                    "experience": []
                }
                print(f"Used regex fallback extraction: {normalized}")
            
            # Ensure all required fields exist
            if 'name' not in normalized or not normalized['name']:
                normalized['name'] = "Unknown Candidate"
            
            if 'email' not in normalized or not normalized['email']:
                normalized['email'] = "no-email@example.com"
            
            # Process the skills field - it might be a list of dicts or a list of strings
            if 'skills' not in normalized:
                normalized['skills'] = []
            else:
                # Check if it's a list or a string
                if isinstance(normalized['skills'], str):
                    normalized['skills'] = [s.strip() for s in normalized['skills'].split(',')]
                
                # If it's a list of objects with 'name' field, flatten it
                if isinstance(normalized['skills'], list) and len(normalized['skills']) > 0 and isinstance(normalized['skills'][0], dict):
                    skills_list = []
                    for skill in normalized['skills']:
                        if isinstance(skill, dict) and 'name' in skill:
                            skills_list.append(skill['name'])
                        elif isinstance(skill, str):
                            skills_list.append(skill)
                    normalized['skills'] = skills_list
                    
                # Extract skills from longer text descriptions
                processed_skills = []
                for skill in normalized['skills']:
                    if isinstance(skill, str):
                        # If the skill is a long text, use an agent-based approach to extract core skills
                        if len(skill) > 100:
                            # Use LLM to extract core technical skills instead of hardcoded list
                            try:
                                tech_skill_prompt = f"Extract specific technical skills from this text. Return only a comma-separated list of skills: {skill}"
                                extracted_skills = self.llm.invoke(tech_skill_prompt).content.split(',')
                                processed_skills.extend([s.strip() for s in extracted_skills if s.strip()])
                            except Exception as e:
                                print(f"Error extracting skills with LLM: {e}")
                                # Add the original if extraction failed
                                processed_skills.append(skill)
                        else:
                            processed_skills.append(skill)
                
                if processed_skills:
                    normalized['skills'] = processed_skills
            
            # Process education
            if 'education' not in normalized:
                normalized['education'] = []
            
            # Process experience
            if 'experience' not in normalized:
                normalized['experience'] = []
                
            print(f"Final normalized candidate info: {normalized}")
            return normalized
            
        except Exception as e:
            print(f"Error extracting candidate info: {e}")
            # Return minimal info
            return {
                "name": "Unknown Candidate",
                "email": "no-email@example.com",
                "skills": [],
                "education": [],
                "experience": []
            }

class MatchingAgent:
    """Agent responsible for matching candidates with job descriptions."""
    
    def __init__(self):
        """Initialize the MatchingAgent."""
        base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        self.llm = ChatOllama(
            model="llama3",
            base_url=base_url,
            temperature=0.1
        )
        self._cache = {}
        # Set threshold for successful match (lowered to match main.py)
        self.match_threshold = 40  # 40% match score for generating interviews
    
    def _get_cache_key(self, *args):
        """Generate a cache key from input arguments"""
        key_str = str(args)
        return hashlib.md5(key_str.encode('utf-8')).hexdigest()
    
    def _normalize_job_requirements(self, job_requirements):
        """Normalize job requirements for consistent matching."""
        
        # Create a copy to avoid modifying the original
        normalized = job_requirements.copy() if job_requirements else {}
        
        # Handle job title extraction more robustly
        if not normalized.get('title') or normalized.get('title') == 'Unknown':
            # Try multiple potential keys for title
            potential_title_keys = ['job_title', 'position', 'role']
            for key in potential_title_keys:
                if key in normalized and normalized[key]:
                    normalized['title'] = normalized[key]
                    break
                    
            # If still not found, try extracting from description
            if (not normalized.get('title') or normalized.get('title') == 'Unknown') and normalized.get('description'):
                # Extract potential title from first few lines of description
                desc_lines = normalized.get('description', '').split('\n')
                for line in desc_lines[:5]:  # Check first 5 lines
                    line = line.strip()
                    # Common patterns for job titles in descriptions
                    if line and len(line) < 100 and any(x in line.lower() for x in ['position:', 'title:', 'role:', 'job:']):
                        # Extract the title after the indicator
                        for indicator in ['position:', 'title:', 'role:', 'job:']:
                            if indicator in line.lower():
                                normalized['title'] = line.split(indicator, 1)[1].strip()
                                break
                        if normalized.get('title'): break

        # Handle required skills more robustly
        if not normalized.get('required_skills'):
            normalized['required_skills'] = []
            
            # Try to find skills in various formats
            skill_keys = ['skills', 'technical_skills', 'technologies', 'tech_stack']
            for key in skill_keys:
                if key in normalized and normalized[key]:
                    skills = normalized[key]
                    if isinstance(skills, str):
                        # Convert string to list
                        skills = [s.strip() for s in skills.split(',')]
                    if isinstance(skills, list):
                        normalized['required_skills'].extend(skills)
                        
            # If still no skills found, try to extract from responsibilities
            if not normalized['required_skills'] and normalized.get('responsibilities'):
                resp = normalized.get('responsibilities', [])
                if isinstance(resp, str):
                    resp = [resp]
                    
                # Use LLM to extract technical skills from responsibilities
                try:
                    resp_text = " ".join(resp)
                    tech_skills_prompt = f"Extract technical skills from these job responsibilities. Return a comma-separated list of only the technical skills: {resp_text}"
                    extracted_skills = self.llm.invoke(tech_skills_prompt).content.split(',')
                    normalized['required_skills'].extend([s.strip() for s in extracted_skills if s.strip()])
                except Exception as e:
                    print(f"Error extracting skills from responsibilities with LLM: {e}")
                    # If LLM fails, use a minimal approach to avoid breaking
                    for item in resp:
                        if "skill" in item.lower() or "technology" in item.lower() or "experience with" in item.lower():
                            potential_skill = item.replace("skill", "").replace("technology", "").replace("experience with", "").strip()
                            if potential_skill and len(potential_skill) < 50:  # Basic sanity check
                                normalized['required_skills'].append(potential_skill)
        
        # Ensure we have lists for structured data
        for key in ['required_skills', 'qualifications', 'responsibilities']:
            if key in normalized and not isinstance(normalized[key], list):
                if isinstance(normalized[key], str):
                    normalized[key] = [s.strip() for s in normalized[key].split(',')]
                else:
                    normalized[key] = []
        
        return normalized
    
    def _calculate_skills_match(self, candidate_skills, job_skills):
        """Calculate the match score for skills."""
        # Normalize skills to lists
        if not isinstance(candidate_skills, list):
            candidate_skills = []
        if not isinstance(job_skills, list):
            job_skills = []
            
        # If empty after normalization, return 0
        if not candidate_skills or not job_skills:
            return 0.0
            
        # Normalize to lowercase for comparison
        candidate_skills_lower = []
        for skill in candidate_skills:
            if isinstance(skill, str):
                candidate_skills_lower.append(skill.lower())
            elif isinstance(skill, dict) and 'name' in skill and isinstance(skill['name'], str):
                candidate_skills_lower.append(skill['name'].lower())
                
        job_skills_lower = []
        for skill in job_skills:
            if isinstance(skill, str):
                job_skills_lower.append(skill.lower())
            elif isinstance(skill, dict) and 'name' in skill and isinstance(skill['name'], str):
                job_skills_lower.append(skill['name'].lower())
        
        print(f"Normalized candidate skills: {candidate_skills_lower}")
        print(f"Normalized job skills: {job_skills_lower}")
        
        # Prepare list for matched skills (for debugging)
        matched_skills = []
        
        # Direct matches
        direct_matches = 0
        for job_skill in job_skills_lower:
            # Skip empty strings and non-string values
            if not isinstance(job_skill, str) or not job_skill.strip():
                continue
                
            # Check for exact matches or if job skill is contained in candidate skill
            if any(isinstance(cs, str) and job_skill == cs for cs in candidate_skills_lower) or any(isinstance(cs, str) and job_skill in cs for cs in candidate_skills_lower):
                direct_matches += 1
                matched_skills.append(job_skill)
                continue
                
            # Try keyword matching using the LLM instead of hardcoded variations
            try:
                # Only check with LLM if we couldn't find a direct match
                candidate_skills_text = ", ".join([cs for cs in candidate_skills_lower if isinstance(cs, str)])
                variation_prompt = f"""
                Does the job skill '{job_skill}' match or closely relate to any of these candidate skills: {candidate_skills_text}?
                Consider common variations, abbreviations, and related technologies.
                Return only 'YES' if there's a match or 'NO' if there isn't.
                """
                variation_result = self.llm.invoke(variation_prompt).content.strip().upper()
                
                if variation_result == "YES":
                    direct_matches += 1
                    matched_skills.append(f"{job_skill} (via semantic match)")
                    continue
            except Exception as e:
                print(f"Error using LLM for skill variation matching: {e}")
                # Fall through to partial matching if LLM fails
        
        # Avoid division by zero
        if not job_skills_lower:
            return 0.0
            
        # Calculate match percentage (count partial matches accordingly)
        score = min(100.0, (direct_matches / len(job_skills_lower)) * 100) if job_skills_lower else 0.0
        
        print(f"Skills Match: {score:.1f}% - Matched {direct_matches}/{len(job_skills_lower)} skills")
        if matched_skills:
            print(f"Matched skills: {matched_skills}")
            
        return score

    def _calculate_qualification_match(self, candidate_qualifications, job_qualifications):
        """Calculate the match score for qualifications."""
        if not candidate_qualifications or not job_qualifications:
            return 0.0

        # Normalize to lists of strings
        candidate_quals_normalized = []
        if isinstance(candidate_qualifications, str):
            candidate_quals_normalized = [q.strip().lower() for q in candidate_qualifications.split(',') if isinstance(q, str)]
        elif isinstance(candidate_qualifications, list):
            for q in candidate_qualifications:
                if isinstance(q, str):
                    candidate_quals_normalized.append(q.strip().lower())
                elif isinstance(q, dict) and 'degree' in q and isinstance(q['degree'], str):
                    # Handle education entries that might be passed in
                    candidate_quals_normalized.append(str(q['degree']).strip().lower())

        job_quals_normalized = []
        if isinstance(job_qualifications, str):
            job_quals_normalized = [q.strip().lower() for q in job_qualifications.split(',') if isinstance(q, str)]
        elif isinstance(job_qualifications, list):
            for q in job_qualifications:
                if isinstance(q, str):
                    job_quals_normalized.append(q.strip().lower())
                elif isinstance(q, dict) and 'degree' in q and isinstance(q['degree'], str):
                    job_quals_normalized.append(str(q['degree']).strip().lower())

        if not candidate_quals_normalized or not job_quals_normalized:
            return 0.0

        # Use LLM to determine education levels instead of hardcoded dictionary
        def get_education_level(education_text):
            try:
                level_prompt = f"""
                Determine the education level value (1-5) for this text: "{education_text}"
                Use this scale:
                1 = High School
                2 = Associate degree
                3 = Bachelor's degree
                4 = Master's degree
                5 = PhD/Doctorate
                
                Return ONLY the numeric value (1-5).
                """
                level_response = self.llm.invoke(level_prompt).content.strip()
                match = re.search(r'(\d+)', level_response)
                if match:
                    return int(match.group(1))
                return 0
            except Exception as e:
                print(f"Error determining education level with LLM: {e}")
                # Fall back to basic keyword matching
                text = education_text.lower()
                if "phd" in text or "doctorate" in text:
                    return 5
                elif "master" in text:
                    return 4
                elif "bachelor" in text or "bs" in text or "ba" in text:
                    return 3
                elif "associate" in text:
                    return 2
                elif "high school" in text:
                    return 1
                return 0

        # Extract required degree level from job qualifications
        required_level = 0
        for qual in job_quals_normalized:
            if not isinstance(qual, str):
                continue
            level_value = get_education_level(qual)
            required_level = max(required_level, level_value)
        
        # Extract highest candidate education level
        candidate_level = 0
        for qual in candidate_quals_normalized:
            if not isinstance(qual, str):
                continue
            level_value = get_education_level(qual)
            candidate_level = max(candidate_level, level_value)
        
        # If no specific education is required, return 1.0
        if required_level == 0:
            return 1.0
            
        # If candidate meets or exceeds the requirement
        if candidate_level >= required_level:
            return 1.0
            
        # Partial credit for some education
        if candidate_level > 0:
            return candidate_level / required_level
            
        return 0.0

    def calculate_match_score(self, candidate_info, job_requirements):
        """
        Calculate the match score between a candidate and a job.

        Args:
            candidate_info: Dictionary containing candidate information
            job_requirements: Dictionary containing job requirements

        Returns:
            Dictionary containing match scores
        """
        print("\nCalculating match score...")

        # First normalize the objects if they have 'properties'
        if isinstance(candidate_info, dict) and 'properties' in candidate_info:
            print("Normalizing candidate info from properties")
            properties = candidate_info.get('properties', {})
            normalized_candidate = {}
            for key, value in properties.items():
                normalized_candidate[key] = value
            # Don't overwrite with empty fields
            for key, value in candidate_info.items():
                if key != 'properties' and key != 'required' and key not in normalized_candidate:
                    normalized_candidate[key] = value
            candidate_info = normalized_candidate

        if isinstance(job_requirements, dict) and 'properties' in job_requirements:
            print("Normalizing job requirements from properties")
            properties = job_requirements.get('properties', {})
            normalized_job = {}
            for key, value in properties.items():
                normalized_job[key] = value
            # Don't overwrite with empty fields
            for key, value in job_requirements.items():
                if key != 'properties' and key != 'required' and key not in normalized_job:
                    normalized_job[key] = value
            job_requirements = normalized_job

        # Create cache key for this calculation
        cache_key = self._get_cache_key(str(candidate_info), str(job_requirements))
        if cache_key in self._cache:
            print("Using cached match score")
            return self._cache[cache_key]

        # Normalize the job requirements and candidate info
        normalized_job = self._normalize_job_requirements(job_requirements)

        print(f"Job Title: {normalized_job.get('title', 'Unknown')}")
        print(f"Required Skills: {normalized_job.get('required_skills', [])}")
        print(f"Required Experience (Years): {normalized_job.get('experience_years', 'N/A')}")
        print(f"Required Qualifications: {normalized_job.get('qualifications', [])}")

        # Extract candidate details
        candidate_skills = candidate_info.get('skills', [])
        candidate_experience = candidate_info.get('experience', [])  # Assume list of dicts with 'years'
        candidate_education = candidate_info.get('education', [])  # Assume list of dicts with 'degree'
        candidate_qualifications = candidate_info.get('qualifications', [])  # Extract candidate qualifications

        # Extract job requirements details
        job_skills = normalized_job.get('required_skills', [])
        
        # Ensure experience_years is a number
        job_experience_years = 0
        try:
            if 'experience_years' in normalized_job:
                # Convert to float first to handle both integers and strings with decimals
                job_experience_years = float(normalized_job['experience_years'])
        except (ValueError, TypeError):
            # If conversion fails, default to 0
            print(f"Could not convert job experience years to number: {normalized_job.get('experience_years')}")
            job_experience_years = 0
            
        job_education = normalized_job.get('education', [])  # Typically part of qualifications
        job_qualifications = normalized_job.get('qualifications', [])

        # --- Calculate individual scores ---

        # 1. Skills Match (Weight: 40%)
        skills_score = self._calculate_skills_match(candidate_skills, job_skills)

        # 2. Experience Match (Weight: 30%)
        # Extract total years of experience from candidate info
        candidate_total_years = 0
        if isinstance(candidate_experience, list):
            for exp in candidate_experience:
                if isinstance(exp, dict) and 'years' in exp:
                    try:
                        # Ensure we're working with numeric values
                        years_value = exp['years']
                        if isinstance(years_value, str):
                            # Try to extract numeric value from strings like "3 years"
                            numeric_match = re.search(r'(\d+(?:\.\d+)?)', years_value)
                            if numeric_match:
                                candidate_total_years += float(numeric_match.group(1))
                        else:
                            candidate_total_years += float(years_value)
                    except (ValueError, TypeError):
                        print(f"Could not convert experience years to number: {exp.get('years')}")
                        pass  # Ignore invalid years format
                # Also check for 'duration' if 'years' is not available
                elif isinstance(exp, dict) and 'duration' in exp and not 'years' in exp:
                    try:
                        # Try to extract years from duration strings like "2 years 3 months"
                        duration_str = exp['duration']
                        years_match = re.search(r'(\d+)\s*(?:years|year)', duration_str, re.IGNORECASE)
                        months_match = re.search(r'(\d+)\s*(?:months|month)', duration_str, re.IGNORECASE)
                        
                        if years_match:
                            candidate_total_years += float(years_match.group(1))
                        if months_match:
                            candidate_total_years += float(months_match.group(1)) / 12
                    except (ValueError, TypeError, AttributeError):
                        pass  # Ignore invalid duration format

        experience_score = 0.0
        if job_experience_years > 0:
            # Simple linear score: 100% if candidate meets or exceeds required years
            # Scales down if candidate has less experience
            ratio = min(1.0, candidate_total_years / job_experience_years)
            experience_score = ratio * 100
        else:
            experience_score = 100.0  # If no experience required, score is 100
        print(f"Experience Match: {experience_score:.1f}% (Cand: {candidate_total_years} yrs, Req: {job_experience_years} yrs)")

        # 3. Education Match (Weight: 15%) - often part of qualifications, simplified check
        # Example: Check if candidate has at least a degree mentioned in job reqs
        education_score = 0.0
        required_degrees = [q.lower() for q in job_qualifications if isinstance(q, str) and any(deg in q.lower() for deg in ["degree", "bachelor", "master", "phd", "bs", "ms"])]
        if not required_degrees:
            education_score = 100.0  # No specific degree required
        else:
            candidate_degrees = []
            if isinstance(candidate_education, list):
                candidate_degrees = [str(edu.get('degree', '')).lower() for edu in candidate_education if isinstance(edu, dict) and edu.get('degree')]

            if any(req_deg in cand_deg for req_deg in required_degrees for cand_deg in candidate_degrees):
                education_score = 100.0  # Candidate has at least one required degree type
            else:
                # Could add partial credit here for related degrees
                education_score = 0.0
        print(f"Education Match: {education_score:.1f}%")

        # 4. Qualifications Match (Weight: 15%)
        # Combine job education requirements into qualifications for matching
        combined_job_quals = []
        # Ensure we only add strings to the combined qualifications list
        if isinstance(job_qualifications, list):
            combined_job_quals.extend([q for q in job_qualifications if isinstance(q, str)])
        if isinstance(job_education, list):
            combined_job_quals.extend([edu for edu in job_education if isinstance(edu, str)])
            
        qualification_score = self._calculate_qualification_match(
            [q for q in candidate_qualifications if isinstance(q, str)] + 
            [str(edu.get('degree', '')) for edu in candidate_education if isinstance(edu, dict) and edu.get('degree')], 
            combined_job_quals
        )

        # --- Calculate Overall Score ---
        # Weights: Skills (0.4), Experience (0.3), Education (0.15), Qualifications (0.15)
        overall_score = (
            (skills_score * 0.40) +
            (experience_score * 0.30) +
            (education_score * 0.15) +
            (qualification_score * 0.15)
        )
        overall_score = round(overall_score, 1)  # Round to one decimal place

        print(f"\nOverall Match Score: {overall_score}%")

        # Check against threshold (ensure threshold is a number)
        threshold = float(self.match_threshold)
        threshold_met = overall_score >= threshold

        result = {
            "overall_score": overall_score,
            "skills_match": skills_score,
            "experience_match": experience_score,
            "education_match": education_score,  # Might be redundant if covered by qualifications
            "qualification_match": qualification_score,
            "threshold_met": threshold_met
        }

        self._cache[cache_key] = result
        return result

    def generate_interview_request(self, candidate: CandidateInfo, job: JobRequirements, match_score: Dict[str, float]) -> Optional[Dict]:
        """
        Generate interview request email content for high-matching candidates.
        
        Args:
            candidate: Candidate information
            job: Job requirements
            match_score: Match score information
            
        Returns:
            Dictionary with interview request details or None if not appropriate
        """
        # Only generate interview request if threshold met
        if not match_score.get('threshold_met', False):
            print("Threshold not met, skipping interview request generation")
            return None
        
        # Get candidate name
        candidate_name = candidate.get('name', 'Candidate')
        job_title = job.get('title', 'the position')
        company_name = job.get('company', 'our company')
        skills_match = match_score.get('skills_match', 0)
        
        # Get candidate skills that matched
        candidate_skills = candidate.get('skills', [])
        job_skills = job.get('required_skills', [])
        matching_skills = []
        
        if candidate_skills and job_skills:
            # Find skills that match between candidate and job
            for skill in candidate_skills:
                if isinstance(skill, str) and any(isinstance(s, str) and skill.lower() in s.lower() for s in job_skills):
                    matching_skills.append(skill)
        
        # Build prompt for interview request generation
        prompt = f"""
        Generate a professional and personalized interview request email for {candidate_name} who has matched well with the {job_title} position at {company_name}.
        The match score was {match_score.get('overall_score', 0):.1f}%, with skills match of {skills_match:.1f}%.
        
        {f"Their skills in {', '.join(matching_skills[:3])} align well with our requirements." if matching_skills else ""}
        
        The email should include:
        1. A personalized greeting addressing {candidate_name} by name
        2. A brief mention of their specific match with the position, mentioning their key skills
        3. Three potential interview dates in the next two weeks (use actual dates)
        4. Whether this will be remote, in-person, or hybrid
        5. A brief mention of next steps in the process
        6. A professional closing with a real name and title
        
        Format the response as a JSON object with fields:
        - subject: the email subject line (personalized with candidate name and job title)
        - body: the complete email body in HTML format (well-formatted with paragraphs)
        - proposed_dates: array of three potential interview dates (as strings in YYYY-MM-DD format)
        - interview_type: either "remote", "in-person", or "hybrid"
        """
        
        try:
            # Generate the interview request
            response = self.llm.invoke([HumanMessage(content=prompt)])
            content = response.content
            
            if isinstance(content, str):
                # Try to parse JSON from the response
                try:
                    json_content = json.loads(content)
                    
                    # Validate response format
                    if not all(k in json_content for k in ['subject', 'body', 'proposed_dates', 'interview_type']):
                        print("Generated interview request missing required fields")
                        return None
                    
                    # Format the body as HTML if it's not already
                    if not json_content['body'].startswith('<'):
                        formatted_body = ''
                        paragraphs = json_content['body'].split('\n\n')
                        for p in paragraphs:
                            if p.strip():
                                formatted_body += f"<p>{p.strip()}</p>\n"
                        json_content['email_body'] = formatted_body
                    else:
                        json_content['email_body'] = json_content['body']
                    
                    return json_content
                except json.JSONDecodeError:
                    print("Failed to parse JSON from interview request generation")
                    return None
            else:
                # Already in parsed JSON format
                return content
        except Exception as e:
            print(f"Error generating interview request: {e}")
            return None

    def match_skills(self, candidate_skills, job_skills):
        """Match candidate skills with job skills."""
        # Handle different data structures
        candidate_skills_lower = []
        for skill in candidate_skills:
            if isinstance(skill, str):
                candidate_skills_lower.append(skill.lower())
            elif isinstance(skill, dict) and 'name' in skill:
                candidate_skills_lower.append(skill['name'].lower())
        
        job_skills_lower = []
        for skill in job_skills:
            if isinstance(skill, str):
                job_skills_lower.append(skill.lower())
            elif isinstance(skill, dict) and 'name' in skill:
                job_skills_lower.append(skill['name'].lower())
        
        # Count matches
        matches = 0
        for c_skill in candidate_skills_lower:
            for j_skill in job_skills_lower:
                if isinstance(c_skill, str) and isinstance(j_skill, str) and j_skill in c_skill:
                    matches += 1
                    break
        
        if not job_skills_lower:
            return 1.0  # No skills required, perfect match
            
        match_score = matches / len(job_skills_lower)
        return min(1.0, match_score)  # Cap at 100%

    def match_qualifications(self, candidate_qualifications, job_qualifications):
        """Match candidate qualifications with job requirements."""
        # Handle different data structures
        candidate_quals_normalized = []
        
        if isinstance(candidate_qualifications, str):
            candidate_quals_normalized = [q.strip().lower() for q in candidate_qualifications.split(',') if isinstance(q, str)]
        elif isinstance(candidate_qualifications, list):
            for q in candidate_qualifications:
                if isinstance(q, str):
                    candidate_quals_normalized.append(q.strip().lower())
                elif isinstance(q, dict) and 'degree' in q:
                    candidate_quals_normalized.append(str(q['degree']).strip().lower())
        
        job_quals_normalized = []
        if isinstance(job_qualifications, str):
            job_quals_normalized = [q.strip().lower() for q in job_qualifications.split(',') if isinstance(q, str)]
        elif isinstance(job_qualifications, list):
            for q in job_qualifications:
                if isinstance(q, str):
                    job_quals_normalized.append(q.strip().lower())
                elif isinstance(q, dict) and 'degree' in q:
                    job_quals_normalized.append(str(q['degree']).strip().lower())
        
        # Use LLM to determine education levels instead of hardcoded dictionary
        def get_education_level(education_text):
            try:
                level_prompt = f"""
                Determine the education level value (1-5) for this text: "{education_text}"
                Use this scale:
                1 = High School
                2 = Associate degree
                3 = Bachelor's degree
                4 = Master's degree
                5 = PhD/Doctorate
                
                Return ONLY the numeric value (1-5).
                """
                level_response = self.llm.invoke(level_prompt).content.strip()
                match = re.search(r'(\d+)', level_response)
                if match:
                    return int(match.group(1))
                return 0
            except Exception as e:
                print(f"Error determining education level with LLM: {e}")
                # Fall back to basic keyword matching
                text = education_text.lower()
                if "phd" in text or "doctorate" in text:
                    return 5
                elif "master" in text:
                    return 4
                elif "bachelor" in text or "bs" in text or "ba" in text:
                    return 3
                elif "associate" in text:
                    return 2
                elif "high school" in text:
                    return 1
                return 0

        # Extract required degree level from job qualifications
        required_level = 0
        for qual in job_quals_normalized:
            if not isinstance(qual, str):
                continue
            level_value = get_education_level(qual)
            required_level = max(required_level, level_value)
        
        # Extract highest candidate education level
        candidate_level = 0
        for qual in candidate_quals_normalized:
            if not isinstance(qual, str):
                continue
            level_value = get_education_level(qual)
            candidate_level = max(candidate_level, level_value)
        
        # If no specific education is required, return 1.0
        if required_level == 0:
            return 1.0
            
        # If candidate meets or exceeds the requirement
        if candidate_level >= required_level:
            return 1.0
            
        # Partial credit for some education
        if candidate_level > 0:
            return candidate_level / required_level
            
        return 0.0